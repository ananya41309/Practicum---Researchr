import os
from PyPDF2 import PdfReader
from docx import Document

def parse_txt(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def parse_pdf(file_path):
    text = ""
    reader = PdfReader(file_path)
    for page in reader.pages:
        if page.extract_text():
            text += page.extract_text() + " "
    return text

def parse_docx(file_path):
    doc = Document(file_path)
    return " ".join([para.text for para in doc.paragraphs])

def parse_file(file_path):
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".txt":
        return parse_txt(file_path)
    elif extension == ".pdf":
        return parse_pdf(file_path)
    elif extension == ".docx":
        return parse_docx(file_path)
    else:
        return ""
import os
from PyPDF2 import PdfReader
from docx import Document

def parse_txt(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def parse_pdf(file_path):
    text = ""
    reader = PdfReader(file_path)
    for page in reader.pages:
        if page.extract_text():
            text += page.extract_text() + " "
    return text

def parse_docx(file_path):
    doc = Document(file_path)
    return " ".join([para.text for para in doc.paragraphs])

def parse_file(file_path):
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".txt":
        return parse_txt(file_path)
    elif extension == ".pdf":
        return parse_pdf(file_path)
    elif extension == ".docx":
        return parse_docx(file_path)
    else:
        return ""
